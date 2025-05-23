import sys
import os
import gc
from pathlib import Path
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QTextEdit, QPushButton, QLabel, 
                            QFileDialog, QMessageBox, QProgressDialog,
                            QStatusBar, QSplashScreen, QScrollArea, QFrame,
                            QStackedWidget, QSizePolicy, QLineEdit)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer, QDateTime, QSize, QRect
from PyQt6.QtGui import QFont, QIcon, QPixmap, QColor, QPainter, QLinearGradient, QPainterPath
from agno.agent import Agent
from agno.models.openai import OpenAIChat
from dotenv import load_dotenv, find_dotenv, set_key
from agno.embedder.openai import OpenAIEmbedder
import logging
import psutil
import time
from PyPDF2 import PdfReader
import json
import uuid
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import re
import docx  # Import for processing .docx files

# Import our persistence module
from persistence import VectorStorePersistence

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Application constants
APP_TITLE = "iATAS - Analisador de ATAS"
APP_VERSION = "1.2.0"
APP_DATA_DIR = "app_data"

class SettingsManager:
    """Handles application settings persistence"""
    
    def __init__(self, app_data_dir=APP_DATA_DIR):
        """Initialize settings manager"""
        self.app_data_dir = Path(app_data_dir)
        self.settings_file = self.app_data_dir / "settings.json"
        self._ensure_dir_exists()
        self.settings = self.load_settings()
        
    def _ensure_dir_exists(self):
        """Create the app data directory if it doesn't exist"""
        self.app_data_dir.mkdir(exist_ok=True, parents=True)
        
    def load_settings(self):
        """Load settings from file"""
        if not self.settings_file.exists():
            return {"openai_api_key": ""}
            
        try:
            with open(self.settings_file, 'r') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Error loading settings: {str(e)}")
            return {"openai_api_key": ""}
            
    def save_settings(self):
        """Save settings to file"""
        try:
            with open(self.settings_file, 'w') as f:
                json.dump(self.settings, f)
            logger.info("Settings saved successfully")
            return True
        except Exception as e:
            logger.error(f"Error saving settings: {str(e)}")
            return False
            
    def get_api_key(self):
        """Get the OpenAI API key"""
        return self.settings.get("openai_api_key", "")
        
    def set_api_key(self, api_key):
        """Set the OpenAI API key"""
        self.settings["openai_api_key"] = api_key
        # Also set the environment variable for immediate use
        os.environ["OPENAI_API_KEY"] = api_key
        return self.save_settings()
        
    def apply_settings(self):
        """Apply settings to the environment"""
        # Set environment variables based on settings
        if "openai_api_key" in self.settings and self.settings["openai_api_key"]:
            os.environ["OPENAI_API_KEY"] = self.settings["openai_api_key"]
            logger.info("Applied API key from settings")

def get_memory_usage():
    """Get current memory usage in MB"""
    process = psutil.Process(os.getpid())
    return process.memory_info().rss / 1024 / 1024

class SimpleVectorStore:
    """Simple vector store implementation using numpy and cosine similarity"""
    def __init__(self, embedder):
        self.embedder = embedder
        self.documents = []
        self.vectors = []
        self.metadatas = []
        self.ids = []

    def add(self, texts, metadatas=None, ids=None):
        """Add documents to the vector store"""
        if metadatas is None:
            metadatas = [{}] * len(texts)
        if ids is None:
            ids = [str(uuid.uuid4()) for _ in texts]

        # Get embeddings for the texts using get_embedding method
        embeddings = []
        for text in texts:
            embeddings.append(self.embedder.get_embedding(text))
        
        # Add to store
        self.documents.extend(texts)
        self.vectors.extend(embeddings)
        self.metadatas.extend(metadatas)
        self.ids.extend(ids)

    def search(self, query, num_documents=5, **kwargs):
        """Search method compatible with agno's expected interface"""
        # Handle None or invalid num_documents parameter
        if num_documents is None or not isinstance(num_documents, int) or num_documents <= 0:
            logger.info(f"Invalid num_documents value: {num_documents}, using default of 5")
            num_documents = 5
            
        results = self.similarity_search(query, k=num_documents)
        
        # Convert to Document objects for agno compatibility
        documents = []
        for result in results:
            # Create a Document object with a to_dict method
            doc = AgnoCompatibleDocument(
                content=result["text"],
                metadata=result["metadata"],
                id=result["id"],
                score=result["score"]
            )
            documents.append(doc)
        
        return documents

    def similarity_search(self, query, k=5):
        """Search for similar documents"""
        if not self.documents:
            return []

        # Handle None k value
        if k is None or not isinstance(k, int) or k <= 0:
            logger.info(f"Invalid k value: {k}, using default of 5")
            k = 5

        # Get query embedding using get_embedding method
        query_embedding = self.embedder.get_embedding(query)
        
        # Calculate similarities
        similarities = cosine_similarity([query_embedding], self.vectors)[0]
        
        # Ensure we don't exceed the number of available documents
        k = min(k, len(self.documents))
        if k == 0:
            return []
            
        # Get top k results
        top_k_idx = np.argsort(similarities)[-k:][::-1]
        
        results = []
        for idx in top_k_idx:
            results.append({
                'text': self.documents[idx],
                'metadata': self.metadatas[idx],
                'id': self.ids[idx],
                'score': float(similarities[idx])
            })
        
        return results

class WorkerThread(QThread):
    """Worker thread for processing AI responses"""
    response_ready = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    def __init__(self, agent, question):
        super().__init__()
        self.agent = agent
        self.question = question
        
    @staticmethod
    def _clean_ansi_and_formatting(text):
        """Remove ANSI color codes, box drawing characters and other formatting artifacts"""
        import re
        
        if not text:
            return ""
            
        # Remove ANSI escape sequences (color codes)
        ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
        result = ansi_escape.sub('', text)
        
        # Remove box drawing and formatting characters (pipes, etc.)
        result = re.sub(r'[│┌┐└┘┐┘─┬┴┼┤├]', '', result)
        
        # Remove any remaining formatting artifacts
        result = re.sub(r'\[3[0-9]m', '', result)  # Color markers like [34m
        result = re.sub(r'\[0m', '', result)       # Reset markers
        
        # Further cleanup of any artifacts
        return result.strip()

    def run(self):
        try:
            # Use the agent interface but capture the response directly
            # The print_response method may be printing to console but not returning the value

            # First, search for relevant context
            if hasattr(self.agent, 'knowledge') and self.agent.search_knowledge:
                logger.info("Searching knowledge base for relevant context")
                docs = self.agent.knowledge.search(self.question)
                context_texts = [d.content for d in docs]
                context = "\n\n".join(context_texts)
                logger.info(f"Found {len(docs)} relevant documents")
            else:
                context = ""
                
            # Use the complete_chat method directly 
            if hasattr(self.agent.model, 'complete_chat'):
                # For OpenAI models, build the messages
                if context:
                    messages = [
                        {"role": "system", "content": self.agent.instructions + "\nIMPORTANT: Provide ONLY your final answer. Do NOT include your thinking process, reasoning, or intermediate steps in your response."},
                        {"role": "user", "content": f"Contexto:\n{context}\n\nPergunta: {self.question}"}
                    ]
                else:
                    messages = [
                        {"role": "system", "content": self.agent.instructions + "\nIMPORTANT: Provide ONLY your final answer. Do NOT include your thinking process, reasoning, or intermediate steps in your response."},
                        {"role": "user", "content": self.question}
                    ]
                
                logger.info("Getting response from model using complete_chat")
                response = self.agent.model.complete_chat(messages)
                logger.info(f"Got response from model: {len(response)} characters")
            else:
                # Fall back to the standard agent method but capture the result
                # Temporarily redirect stdout to capture the output
                import io
                from contextlib import redirect_stdout
                
                f = io.StringIO()
                with redirect_stdout(f):
                    self.agent.print_response(self.question, markdown=True)
                
                output = f.getvalue()
                
                # Extract just the actual response text from the captured output
                # Split the output to remove headers/footers added by agno framework
                lines = output.split("\n")
                
                # Skip past headers to get to the response
                response_lines = []
                collecting = False
                found_response = False
                thinking_mode = False
                
                for line in lines:
                    # Skip sections that indicate thinking process
                    if any(marker in line.lower() for marker in ["thinking", "thought process", "reasoning", "let me think", "analyzing"]):
                        thinking_mode = True
                        continue
                        
                    # Exit thinking mode when we find a conclusion/result marker
                    if thinking_mode and any(marker in line.lower() for marker in ["answer:", "response:", "conclusion:", "therefore,"]):
                        thinking_mode = False
                    
                    # Start collecting after we see the Response header
                    if "Response" in line and "──────" in line:
                        collecting = True
                        found_response = True
                        continue
                        
                    # Stop collecting when we hit the end marker
                    if collecting and "───────────────────" in line:
                        collecting = False
                        continue
                        
                    # Collect the actual response text, skipping thinking sections
                    if collecting and not thinking_mode:
                        # Remove any leading/trailing whitespace and artifacts
                        clean_line = line.strip()
                        # Strip ANSI color codes, border characters, and other artifacts
                        clean_line = self._clean_ansi_and_formatting(clean_line)
                        if clean_line:  # Skip empty lines
                            response_lines.append(clean_line)
                
                # If we couldn't parse the output, just use it as is
                if not found_response or not response_lines:
                    logger.warning("Could not extract clean response text, using raw output")
                    response = self._clean_ansi_and_formatting(output)
                    
                    # Final cleanup to remove thinking process markers and sections
                    response_parts = response.split("\n")
                    final_parts = []
                    skip_section = False
                    
                    for part in response_parts:
                        if any(marker in part.lower() for marker in ["thinking", "thought process", "reasoning", "let me think", "analyzing"]):
                            skip_section = True
                            continue
                            
                        if skip_section and any(marker in part.lower() for marker in ["answer:", "response:", "conclusion:", "therefore,"]):
                            skip_section = False
                            # Keep this line as it likely has the answer
                            part = part.replace("Answer:", "").replace("Response:", "").replace("Conclusion:", "").strip()
                            final_parts.append(part)
                            continue
                            
                        if not skip_section:
                            final_parts.append(part)
                    
                    response = "\n".join(final_parts)
                else:
                    # Join the cleaned response lines
                    response = "\n".join(response_lines)
                
                logger.info(f"Processed response text (length: {len(response)} characters)")
            
            if not response or len(response.strip()) < 5:
                raise ValueError("Resposta vazia ou muito curta recebida do modelo")
            
            # Final cleanup to ensure no "thinking" sections remain
            response_parts = response.split("\n")
            final_response = []
            in_thinking_section = False
            
            for part in response_parts:
                # Check for thinking section headers
                if any(marker in part.lower() for marker in ["thinking:", "thought process:", "reasoning:", "let me think", "analyzing"]):
                    in_thinking_section = True
                    continue
                
                # Check for end of thinking section
                if in_thinking_section and any(marker in part.lower() for marker in ["answer:", "response:", "conclusion:", "therefore,"]):
                    in_thinking_section = False
                    # Extract just the answer part
                    for prefix in ["answer:", "response:", "conclusion:", "therefore,"]:
                        if prefix in part.lower():
                            part = part[part.lower().find(prefix) + len(prefix):].strip()
                            break
                
                # Add lines that aren't in thinking sections
                if not in_thinking_section:
                    final_response.append(part)
            
            response = "\n".join(final_response)
            self.response_ready.emit(response)
            
        except Exception as e:
            logger.error(f"Error processing question: {str(e)}")
            self.error_occurred.emit(str(e))

class DocumentLoaderThread(QThread):
    """Worker thread for loading documents"""
    progress = pyqtSignal(int)
    finished = pyqtSignal()
    error = pyqtSignal(str)
    file_processed = pyqtSignal(str)
    memory_warning = pyqtSignal(str)
    status_update = pyqtSignal(str)

    def __init__(self, vector_db, persistence=None):
        super().__init__()
        self.vector_db = vector_db
        self.persistence = persistence
        self.folder_path = None
        self._stop_requested = False

    def set_folder_path(self, path):
        """Set the folder path for document processing"""
        self.folder_path = path

    def stop(self):
        """Request the thread to stop"""
        self._stop_requested = True

    def run(self):
        """Process documents from the selected folder"""
        if not self.folder_path:
            self.error.emit("Nenhuma pasta selecionada")
            return

        try:
            # Get all PDF and TXT files in the folder
            files = []
            for ext in ['*.pdf', '*.txt', '*.docx', '*.doc']:
                files.extend(Path(self.folder_path).glob(ext))
            
            if not files:
                self.error.emit("Nenhum documento encontrado na pasta selecionada")
                return

            total_files = len(files)
            processed_files = 0

            # Clear existing documents before processing new ones
            self.vector_db.documents = []
            self.vector_db.vectors = []
            self.vector_db.metadatas = []
            self.vector_db.ids = []

            for file_path in files:
                if self._stop_requested:
                    break

                try:
                    self.status_update.emit(f"Processando {file_path.name}...")
                    
                    # Extract text based on file type
                    file_suffix = file_path.suffix.lower()
                    if file_suffix == '.pdf':
                        text = self.extract_text_from_pdf(str(file_path))
                    elif file_suffix in ['.docx', '.doc']:
                        text = self.extract_text_from_word(str(file_path))
                    elif file_suffix == '.txt':
                        text = self.extract_text_from_txt(str(file_path))
                    else:
                        # Skip files with unsupported formats
                        self.status_update.emit(f"Aviso: Formato não suportado: {file_path.name}")
                        logger.warning(f"Skipping unsupported file format: {file_path}")
                        continue

                    if not text:
                        self.status_update.emit(f"Aviso: Nenhum texto extraído de {file_path.name}")
                        continue

                    # Process text into chunks
                    chunks = self.process_text(text)
                    
                    if not chunks:
                        self.status_update.emit(f"Aviso: Nenhum chunk gerado de {file_path.name}")
                        continue

                    # Create metadata for each chunk
                    metadatas = []
                    for i, chunk in enumerate(chunks):
                        metadata = {
                            'source': str(file_path),
                            'chunk_index': i,
                            'total_chunks': len(chunks),
                            'file_name': file_path.name,
                            'file_type': file_path.suffix.lower()[1:],
                            'chunk_id': str(uuid.uuid4())
                        }
                        metadatas.append(metadata)

                    # Add to vector store with metadata
                    self.vector_db.add(
                        texts=chunks,
                        metadatas=metadatas
                    )
                    
                    processed_files += 1
                    progress = int((processed_files / total_files) * 100)
                    self.progress.emit(progress)
                    self.file_processed.emit(f"Processado: {file_path.name}")
                    
                    # Check memory usage
                    memory_usage = get_memory_usage()
                    if memory_usage > 1000:  # Warning at 1GB
                        self.memory_warning.emit(
                            f"Uso de memória alto ({memory_usage:.1f}MB). "
                            "Considere processar menos documentos por vez."
                        )

                except Exception as e:
                    logger.error(f"Error processing file {file_path}: {str(e)}")
                    self.error.emit(f"Erro ao processar {file_path.name}: {str(e)}")
                    continue

            if not self._stop_requested:
                # Save the vector store after processing all files
                if hasattr(self, 'persistence') and self.persistence and len(self.vector_db.documents) > 0:
                    try:
                        self.persistence.save_vector_store(self.vector_db)
                        logger.info("Vector store saved successfully")
                    except Exception as e:
                        logger.error(f"Error saving vector store: {str(e)}")
                        self.error.emit(f"Erro ao salvar dados: {str(e)}")
                self.finished.emit()

        except Exception as e:
            logger.error(f"Error in document loader thread: {str(e)}")
            self.error.emit(f"Erro ao processar documentos: {str(e)}")

    def extract_text_from_pdf(self, pdf_path):
        """Extract text from PDF file"""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                total_pages = len(reader.pages)
                
                # Show progress even inside the PDF processing
                for i, page in enumerate(reader.pages):
                    if self._stop_requested:
                        break
                    text += page.extract_text() + "\n"
                    
                    # Update status for larger PDFs
                    if total_pages > 10 and i % 5 == 0:
                        self.status_update.emit(f"Processando {pdf_path.name}: página {i+1}/{total_pages}")
                        
                return text
        except Exception as e:
            logger.error(f"Error extracting text from {pdf_path}: {str(e)}")
            raise

    def extract_text_from_txt(self, txt_path):
        """Extract text from text file"""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error extracting text from {txt_path}: {str(e)}")
            raise
            
    def extract_text_from_word(self, word_path):
        """Extract text from Word documents (.doc or .docx)"""
        try:
            # For .docx files, use python-docx
            if word_path.lower().endswith('.docx'):
                doc = docx.Document(word_path)
                # Extract text from paragraphs
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                # Extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                
                return text
            # For .doc files (old format)
            elif word_path.lower().endswith('.doc'):
                # Attempt to use docx for .doc files, it might work for some
                try:
                    doc = docx.Document(word_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text
                except Exception as doc_err:
                    logger.warning(f"Could not process .doc file with docx: {str(doc_err)}")
                    # Inform user that .doc files might not be fully supported
                    self.status_update.emit(f"Aviso: Arquivos .doc antigos podem não ser processados corretamente: {word_path}")
                    return f"[Não foi possível extrair texto completo de {os.path.basename(word_path)}. Considere converter para .docx ou .pdf.]"
        except Exception as e:
            logger.error(f"Error extracting text from Word file {word_path}: {str(e)}")
            raise

    def process_text(self, text, chunk_size=1000, overlap=100):
        """
        Split text into chunks with overlap to maintain context
        
        Args:
            text: Text to process
            chunk_size: Target size for each chunk
            overlap: Number of words to overlap between chunks
        """
        if not text or not text.strip():
            return []

        words = text.split()
        total_words = len(words)
        chunks = []
        
        if total_words == 0:
            return chunks
            
        # For very short texts, just return as is
        if total_words < chunk_size:
            chunks.append(" ".join(words))
            return chunks
            
        # Process into overlapping chunks
        pos = 0
        while pos < total_words:
            end_pos = min(pos + chunk_size, total_words)
            chunk = " ".join(words[pos:end_pos])
            chunks.append(chunk)
            
            # Move position forward, accounting for overlap
            pos += chunk_size - overlap
            
            # Avoid getting stuck at the end with too small chunks
            if end_pos == total_words and len(chunks) > 1:
                break
                
            # Show progress for very large documents
            if total_words > 10000 and len(chunks) % 10 == 0:
                progress = int(end_pos / total_words * 100)
                self.status_update.emit(f"Dividindo texto em partes: {progress}%")

        return chunks

    def _on_documents_loaded(self):
        """Handle successful document loading"""
        self.progress_dialog.close()
        self.load_docs_button.setEnabled(True)
        
        # Check for API key
        has_api_key = hasattr(self, 'settings_manager') and self.settings_manager.get_api_key() != ""
        
        # Count unique file sources in metadata instead of total chunks
        sources = set()
        for metadata in self.vector_db.metadatas:
            if "source" in metadata:
                sources.add(metadata["source"])
        
        file_count = len(sources)
        chunk_count = len(self.vector_db.documents)
        
        # Update status - show file count, not chunk count
        self._update_document_status(f"{file_count} documentos carregados")
        self.doc_count_label.setText(f"Documentos: {file_count}")
        
        # Show success message
        QMessageBox.information(self, "Sucesso", f"Documentos carregados com sucesso! ({file_count} arquivos, {chunk_count} fragmentos)")
        self.status_bar.showMessage("Documentos carregados com sucesso!", 5000)
        
        # Enable chat functionality only if we have API key
        if has_api_key:
            self.ask_button.setEnabled(True)
            self.question_input.setEnabled(True)
            self.question_input.setPlaceholderText("Digite sua pergunta sobre os documentos aqui...")
        else:
            self.ask_button.setEnabled(False)
            self.question_input.setEnabled(True)
            self.question_input.setPlaceholderText("Configure sua chave API OpenAI nas configurações para fazer perguntas...")
            self.ask_button.setToolTip("Chave API OpenAI não configurada. Acesse as configurações.")
            
            # Show a message about missing API key
            QMessageBox.warning(
                self, 
                "Chave API Necessária", 
                "Seus documentos foram carregados, mas uma chave API OpenAI é necessária para fazer perguntas.\n\n"
                "Por favor, configure sua chave API nas configurações do aplicativo."
            )
        
        # Save the vector store to disk
        self.persistence.save_vector_store(self.vector_db)

    def _on_document_load_error(self, error_message, progress=None):
        """Handle document loading errors"""
        if progress:
            progress.close()
        self.load_docs_button.setEnabled(True)
        self.ask_button.setEnabled(True)
        self._update_document_status("Erro ao carregar documentos")
        QMessageBox.critical(self, "Erro", f"Erro ao carregar documentos: {error_message}")
        self.status_bar.showMessage(f"Erro: {error_message}", 5000)
        
        
    def closeEvent(self, event):
        """Handle the window close event"""
        # Save persistence data
        try:
            self.status_bar.showMessage("Salvando dados...")
            
            # Save vector store
            if hasattr(self, 'vector_db') and self.vector_db and len(self.vector_db.documents) > 0:
                self.persistence.save_vector_store(self.vector_db)
                
            # Save chat history
            if hasattr(self, 'chat_history') and self.chat_history and len(self.chat_history.history) > 0:
                chat_history_file = os.path.join(APP_DATA_DIR, "chat_history.json")
                self.chat_history.save_to_disk(chat_history_file)
                
        except Exception as e:
            logger.error(f"Error saving data on close: {str(e)}")
            QMessageBox.warning(self, "Erro ao salvar", f"Erro ao salvar dados: {str(e)}")
            
        # Accept the close event
        event.accept()

class SettingsView(QWidget):
    """Settings view for the application"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.initUI()
        # Set white background for the settings view
        self.setStyleSheet("""
            QWidget#settingsView {
                background-color: white;
            }
            QScrollArea, QScrollArea > QWidget > QWidget {
                background-color: white;
            }
            QWidget.card {
                background-color: #f8f9fa;
                border-radius: 8px;
                border: 1px solid #e0e4e8;
            }
            QWidget.card:hover {
                border: 1px solid #c0c4c8;
                box-shadow: 0 4px 8px rgba(0, 0, 0, 0.05);
            }
            QPushButton.back-button {
                background-color: #f0f2f5;
                color: #515769;
                border: 1px solid #e0e4e8;
                padding: 8px 16px;
                border-radius: 8px;
                font-weight: bold;
            }
            QPushButton.back-button:hover {
                background-color: #e9ecf2;
                border: 1px solid #d0d4d8;
            }
            QPushButton.primary-button {
                background-color: #4285f4;
                color: white;
                border: none;
                padding: 10px 16px;
                border-radius: 8px;
                font-weight: bold;
            }
            QPushButton.primary-button:hover {
                background-color: #3367d6;
            }
            QPushButton.primary-button:pressed {
                background-color: #2a56c6;
            }
            QPushButton.danger-button {
                background-color: #f44336;
                color: white;
                border: none;
                padding: 10px 16px;
                border-radius: 8px;
                font-weight: bold;
            }
            QPushButton.danger-button:hover {
                background-color: #e53935;
            }
            QPushButton.danger-button:pressed {
                background-color: #d32f2f;
            }
            QLineEdit.api-input {
                border: 1px solid #dadce0;
                border-radius: 8px;
                padding: 12px;
                background-color: white;
                color: #3c4043;
                font-size: 14px;
            }
            QLineEdit.api-input:focus {
                border: 2px solid #4285f4;
                padding: 11px;
            }
            QLabel.section-title {
                color: #202124;
                font-size: 18px;
                font-weight: bold;
                margin-bottom: 4px;
            }
            QLabel.section-subtitle {
                color: #5f6368;
                font-size: 13px;
                font-weight: normal;
            }
            QLabel.field-label {
                color: #3c4043;
                font-size: 14px;
                font-weight: bold;
            }
            QLabel.version-label {
                color: #80868b;
                font-size: 12px;
            }
        """)
        
    def initUI(self):
        """Initialize the UI components"""
        # Set object name for styling
        self.setObjectName("settingsView")
        
        # Main layout with margins
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)
        
        # Header with title on left and back button on right
        header_layout = QHBoxLayout()
        header_layout.setContentsMargins(0, 0, 0, 10)
        
        # Title on left with icon
        title_layout = QVBoxLayout()
        title_layout.setSpacing(2)
        
        settings_title = QLabel("Configurações")
        settings_title.setFont(QFont("Segoe UI", 20, QFont.Weight.Bold))
        settings_title.setStyleSheet("color: #202124; margin-bottom: 0px;")
        title_layout.addWidget(settings_title)
        
        settings_subtitle = QLabel("Personalize sua experiência com o analisador de ATAS")
        settings_subtitle.setFont(QFont("Segoe UI", 12))
        settings_subtitle.setStyleSheet("color: #5f6368; font-weight: normal;")
        title_layout.addWidget(settings_subtitle)
        
        header_layout.addLayout(title_layout)
        
        # Add stretching space to push button to the right
        header_layout.addStretch()
        
        # Back button on right with icon
        self.back_button = QPushButton(" Voltar")
        self.back_button.setFont(QFont("Segoe UI", 11))
        self.back_button.setProperty("class", "back-button")
        self.back_button.setFixedSize(100, 36)
        self.back_button.setIcon(QIcon.fromTheme("go-previous"))
        self.back_button.clicked.connect(self.go_back)
        header_layout.addWidget(self.back_button, alignment=Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignTop)
        
        layout.addLayout(header_layout)
        
        # Create a scroll area for settings content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        scroll_area.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        
        # Create a container widget for the scroll area
        settings_container = QWidget()
        settings_container.setObjectName("settingsContainer")
        settings_content_layout = QVBoxLayout(settings_container)
        settings_content_layout.setContentsMargins(0, 0, 0, 0)
        settings_content_layout.setSpacing(12)
        
        # API Key Card
        api_card = QWidget()
        api_card.setProperty("class", "card")
        api_card_layout = QVBoxLayout(api_card)
        api_card_layout.setContentsMargins(16, 16, 16, 16)
        api_card_layout.setSpacing(8)
        
        # API Key section header
        api_title = QLabel("Chave API OpenAI")
        api_title.setProperty("class", "section-title")
        api_card_layout.addWidget(api_title)
        
        api_info = QLabel("Configure sua chave de API da OpenAI para usar o aplicativo.")
        api_info.setProperty("class", "section-subtitle")
        api_info.setWordWrap(True)
        api_card_layout.addWidget(api_info)
        
        # Separator line
        api_separator = QFrame()
        api_separator.setFrameShape(QFrame.Shape.HLine)
        api_separator.setFrameShadow(QFrame.Shadow.Sunken)
        api_separator.setStyleSheet("background-color: #e0e4e8; border: none; height: 1px; margin: 4px 0;")
        api_card_layout.addWidget(api_separator)
        
        # API Key input field with label
        api_key_label = QLabel("Chave de API:")
        api_key_label.setProperty("class", "field-label")
        api_card_layout.addWidget(api_key_label)
        
        self.api_key_input = QLineEdit()
        self.api_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.api_key_input.setProperty("class", "api-input")
        self.api_key_input.setMinimumHeight(38)
        self.api_key_input.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.api_key_input.setPlaceholderText("Digite sua chave API da OpenAI")
        
        # Load the API key if we have it
        if self.parent and hasattr(self.parent, 'settings_manager'):
            api_key = self.parent.settings_manager.get_api_key()
            if api_key:
                self.api_key_input.setText(api_key)
                self.api_key_input.setPlaceholderText("*" * len(api_key))
        
        api_card_layout.addWidget(self.api_key_input)
        
        # Info text about API key
        api_key_info = QLabel("Sua chave API é armazenada localmente e nunca é compartilhada.")
        api_key_info.setStyleSheet("color: #80868b; font-size: 12px; font-style: italic;")
        api_card_layout.addWidget(api_key_info)
        
        # Save API Key button
        self.save_api_key_button = QPushButton("Salvar Chave API")
        self.save_api_key_button.setProperty("class", "primary-button")
        self.save_api_key_button.setFont(QFont("Segoe UI", 12))
        self.save_api_key_button.setMinimumHeight(38)
        self.save_api_key_button.clicked.connect(self.save_api_key)
        api_card_layout.addWidget(self.save_api_key_button)
        
        # Add card to layout
        settings_content_layout.addWidget(api_card)
        
        # Reset Data Card
        reset_card = QWidget()
        reset_card.setProperty("class", "card")
        reset_card_layout = QVBoxLayout(reset_card)
        reset_card_layout.setContentsMargins(16, 16, 16, 16)
        reset_card_layout.setSpacing(8)
        
        # Reset section header
        reset_title = QLabel("Resetar Dados")
        reset_title.setProperty("class", "section-title")
        reset_card_layout.addWidget(reset_title)
        
        reset_info = QLabel("Reset completo da base de documentos. Esta ação não pode ser desfeita.")
        reset_info.setProperty("class", "section-subtitle")
        reset_info.setWordWrap(True)
        reset_card_layout.addWidget(reset_info)
        
        # Separator line
        reset_separator = QFrame()
        reset_separator.setFrameShape(QFrame.Shape.HLine)
        reset_separator.setFrameShadow(QFrame.Shadow.Sunken)
        reset_separator.setStyleSheet("background-color: #e0e4e8; border: none; height: 1px; margin: 4px 0;")
        reset_card_layout.addWidget(reset_separator)
        
        # Warning about reset
        reset_warning = QLabel("⚠️ Todos os documentos e análises serão removidos permanentemente.")
        reset_warning.setStyleSheet("color: #f44336; font-size: 13px; font-weight: bold;")
        reset_warning.setWordWrap(True)
        reset_card_layout.addWidget(reset_warning)
        
        # Reset button (full width)
        self.reset_docs_button = QPushButton("Resetar Documentos")
        self.reset_docs_button.setProperty("class", "danger-button")
        self.reset_docs_button.setFont(QFont("Segoe UI", 12, QFont.Weight.Bold))
        self.reset_docs_button.setMinimumHeight(38)
        self.reset_docs_button.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        
        # Connect to the parent's reset_documents method
        if self.parent and hasattr(self.parent, 'reset_documents'):
            self.reset_docs_button.clicked.connect(self.parent.reset_documents)
        else:
            self.reset_docs_button.setEnabled(False)
            logger.warning("Parent doesn't have reset_documents method, reset button disabled")
            
        reset_card_layout.addWidget(self.reset_docs_button)
        
        # Add card to layout
        settings_content_layout.addWidget(reset_card)
        
        # About Card
        about_card = QWidget()
        about_card.setProperty("class", "card")
        about_card_layout = QVBoxLayout(about_card)
        about_card_layout.setContentsMargins(16, 16, 16, 16)
        about_card_layout.setSpacing(6)
        
        # About section header
        about_title = QLabel("Sobre o Aplicativo")
        about_title.setProperty("class", "section-title")
        about_card_layout.addWidget(about_title)
        
        # App info with version
        app_info = QLabel(f"iATAS - Analisador de ATAS v{APP_VERSION}")
        app_info.setStyleSheet("color: #3c4043; font-size: 14px; font-weight: bold;")
        about_card_layout.addWidget(app_info)
        
        # App description
        app_description = QLabel("Uma ferramenta inteligente para análise e consulta de documentos de ATAS.")
        app_description.setStyleSheet("color: #5f6368; font-size: 13px;")
        app_description.setWordWrap(True)
        about_card_layout.addWidget(app_description)
        
        # Copyright info
        copyright_info = QLabel("© 2023-2024 Todos os direitos reservados")
        copyright_info.setStyleSheet("color: #80868b; font-size: 12px; margin-top: 4px;")
        about_card_layout.addWidget(copyright_info)
        
        # Add card to layout
        settings_content_layout.addWidget(about_card)
        
        # No stretching space at the bottom to ensure all content is visible
        
        # Set the container widget as the scroll area's widget
        scroll_area.setWidget(settings_container)
        
        # Add the scroll area to the main layout with proper weight
        layout.addWidget(scroll_area, 1)  # Give scroll area a stretch factor of 1 to fill available space
        
    def save_api_key(self):
        """Save the API key from the input field"""
        if not self.parent or not hasattr(self.parent, 'settings_manager'):
            QMessageBox.warning(self, "Erro", "Não foi possível salvar a chave API.")
            return
            
        api_key = self.api_key_input.text().strip()
        
        # Handle different cases for the API key
        if not api_key:
            # Ask for confirmation if clearing the API key
            if self.parent.settings_manager.get_api_key():
                confirm = QMessageBox.question(
                    self,
                    "Remover Chave API",
                    "Tem certeza que deseja remover a chave API existente? Não será possível fazer perguntas sem uma chave API.",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                    QMessageBox.StandardButton.No
                )
                
                if confirm == QMessageBox.StandardButton.No:
                    return
            else:
                QMessageBox.warning(self, "Aviso", "Por favor, insira uma chave API válida.")
                return
        else:
            # Basic validation for OpenAI API key format (should start with "sk-")
            if not api_key.startswith("sk-") or len(api_key) < 20:
                confirm = QMessageBox.question(
                    self,
                    "Formato de Chave Suspeito",
                    "A chave API fornecida não parece estar no formato correto da OpenAI (deve começar com 'sk-').\n\n"
                    "Tem certeza que deseja salvar esta chave?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                    QMessageBox.StandardButton.No
                )
                
                if confirm == QMessageBox.StandardButton.No:
                    return
        
        # Save the API key
        success = self.parent.settings_manager.set_api_key(api_key)
        
        if success:
            # Show success message
            QMessageBox.information(self, "Sucesso", "Chave API salva com sucesso!")
            
            # If we have a restart_ai_components method on the parent, call it
            if hasattr(self.parent, 'restart_ai_components'):
                self.parent.restart_ai_components()
                
            # Update button states in the main window
            if hasattr(self.parent, '_ensure_buttons_enabled'):
                self.parent._ensure_buttons_enabled()
        else:
            QMessageBox.warning(self, "Erro", "Não foi possível salvar a chave API. Verifique as permissões de arquivo.")
        
    def go_back(self):
        """Return to the main view"""
        if self.parent and hasattr(self.parent, 'stacked_widget'):
            # Switch to main view (index 1)
            self.parent.stacked_widget.setCurrentIndex(1)
            logger.info("Returning from settings to main view")
            
            # Force update of button states when returning to main view
            if hasattr(self.parent, '_ensure_buttons_enabled'):
                # Allow UI to update first
                QTimer.singleShot(100, self.parent._ensure_buttons_enabled)
                logger.info("Scheduled button state update after returning to main view")

class InitialSetupView(QWidget):
    """Initial setup view for selecting document folder"""
    setup_complete = pyqtSignal(str, str)  # Signal emitted with folder path and API key

    def __init__(self, parent=None):
        super().__init__(parent)
        self.initUI()
        
    def initUI(self):
        # Set white background
        self.setStyleSheet("""
            QWidget {
                background-color: white;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
                border: none;
            }
        """)
        
        # Set minimum window size to ensure content has space to display
        self.setMinimumSize(550, 400)
        
        # Main layout with balanced margins for screen utilization
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(30, 20, 30, 20)  # Increased horizontal margins
        main_layout.setSpacing(15)  # Increased spacing
        
        # Header section with gradient background - fill width but stay compact height
        header = QWidget()
        header.setFixedHeight(70)
        header.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        header.setStyleSheet("""
            QWidget {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #4a6fc3, stop:1 #5c85d6);
                border-radius: 6px;
                border: none;
            }
        """)
        
        header_layout = QVBoxLayout(header)
        header_layout.setContentsMargins(20, 8, 20, 8)
        header_layout.setSpacing(2)
        
        # App title - much smaller
        app_title = QLabel("Bem-vindo ao iATAS")
        app_title.setStyleSheet("""
            color: white;
            font-size: 18px;
            font-weight: bold;
            border: none;
        """)
        header_layout.addWidget(app_title)
        
        # App subtitle - smaller and combined with setup instruction
        app_subtitle = QLabel("Analisador Inteligente de ATAS - Configure para começar")
        app_subtitle.setStyleSheet("""
            color: rgba(255, 255, 255, 0.9);
            font-size: 12px;
            border: none;
        """)
        header_layout.addWidget(app_subtitle)
        
        main_layout.addWidget(header)
        
        # Content Cards Container - ensure it expands to fill space
        content_container = QWidget()
        content_container.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        content_container.setStyleSheet("""
            QWidget {
                background-color: white;
                border: none;
            }
        """)
        
        content_layout = QVBoxLayout(content_container)
        content_layout.setSpacing(15)  # Increased spacing between cards
        content_layout.setContentsMargins(0, 10, 0, 10)  # Adjusted margins
        
        # --------- FOLDER SELECTION CARD ---------
        folder_card = QWidget()
        folder_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        folder_card.setStyleSheet("""
            QWidget {
                background-color: #f8f9fa;
                border-radius: 6px;
                border: none;
            }
        """)
        
        folder_layout = QVBoxLayout(folder_card)
        folder_layout.setContentsMargins(15, 15, 15, 15)  # Increased padding
        folder_layout.setSpacing(8)  # Increased spacing
        
        # Card title with icon - more compact
        folder_title_layout = QHBoxLayout()
        folder_title_layout.setSpacing(8)
        
        folder_icon_label = QLabel("📁")
        folder_icon_label.setStyleSheet("""
            font-size: 14px;
            color: #202124;
            border: none;
        """)
        folder_title_layout.addWidget(folder_icon_label)
        
        folder_title = QLabel("Selecione a Pasta de Documentos")
        folder_title.setStyleSheet("""
            font-size: 13px;
            font-weight: bold;
            color: #202124;
        """)
        folder_title_layout.addWidget(folder_title)
        folder_title_layout.addStretch()
        
        folder_layout.addLayout(folder_title_layout)
        
        # Folder selection button and display
        folder_input_layout = QHBoxLayout()
        folder_input_layout.setSpacing(10)
        
        self.selected_folder_display = QLineEdit()
        self.selected_folder_display.setReadOnly(True)
        self.selected_folder_display.setPlaceholderText("Nenhuma pasta selecionada")
        self.selected_folder_display.setFixedHeight(30)  # Slightly taller for better visibility
        self.selected_folder_display.setStyleSheet("""
            QLineEdit {
                border: 1px solid #dadce0;
                border-radius: 4px;
                padding: 4px 10px;
                background-color: white;
                color: #3c4043;
                font-size: 12px;
            }
            QLineEdit:focus {
                border: 1px solid #4285f4;
            }
        """)
        folder_input_layout.addWidget(self.selected_folder_display, 1)
        
        self.select_folder_btn = QPushButton("Selecionar")
        self.select_folder_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.select_folder_btn.setFixedHeight(30)  # Slightly taller for better visibility
        self.select_folder_btn.setFixedWidth(100)  # Fixed width for better proportion
        self.select_folder_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 0 12px;
                font-size: 12px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #43A047;
            }
            QPushButton:pressed {
                background-color: #388E3C;
            }
        """)
        self.select_folder_btn.clicked.connect(self.select_folder)
        folder_input_layout.addWidget(self.select_folder_btn)
        
        folder_layout.addLayout(folder_input_layout)
        
        # Add stretch between cards for flexible spacing
        content_layout.addWidget(folder_card)
        content_layout.addSpacing(5)  # Add small space between cards
        
        # --------- API KEY CARD ---------
        api_card = QWidget()
        api_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        api_card.setStyleSheet("""
            QWidget {
                background-color: #f8f9fa;
                border-radius: 6px;
                border: none;
            }
        """)
        
        api_layout = QVBoxLayout(api_card)
        api_layout.setContentsMargins(15, 15, 15, 15)  # Increased padding
        api_layout.setSpacing(8)  # Increased spacing
        
        # Card title with icon
        api_title_layout = QHBoxLayout()
        api_title_layout.setSpacing(8)
        
        api_icon_label = QLabel("🔑")
        api_icon_label.setStyleSheet("""
            font-size: 14px;
            color: #202124;
            border: none;
        """)
        api_title_layout.addWidget(api_icon_label)
        
        api_title = QLabel("Chave API OpenAI")
        api_title.setStyleSheet("""
            font-size: 13px;
            font-weight: bold;
            color: #202124;
            border: none;
        """)
        api_title_layout.addWidget(api_title)
        api_title_layout.addStretch()
        
        api_layout.addLayout(api_title_layout)
        
        # API description - more concise
        api_description = QLabel("A chave API é armazenada localmente e nunca compartilhada.")
        api_description.setStyleSheet("""
            color: #5f6368;
            font-size: 11px;
            border: none;
        """)
        api_description.setWordWrap(True)
        api_layout.addWidget(api_description)
        
        # API key input
        self.api_key_input = QLineEdit()
        self.api_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.api_key_input.setPlaceholderText("Digite sua chave API OpenAI (sk-...)")
        self.api_key_input.setFixedHeight(30)  # Slightly taller for better visibility
        self.api_key_input.setStyleSheet("""
            QLineEdit {
                border: 1px solid #dadce0;
                border-radius: 4px;
                padding: 4px 10px;
                background-color: white;
                color: #3c4043;
                font-size: 12px;
            }
            QLineEdit:focus {
                border: 1px solid #4285f4;
            }
        """)
        api_layout.addWidget(self.api_key_input)
        
        # Add info text - more compact
        api_info = QLabel("Obtenha em: platform.openai.com/api-keys")
        api_info.setStyleSheet("""
            color: #80868b;
            font-size: 10px;
            font-style: italic;
            border: none;
        """)
        api_info.setOpenExternalLinks(True)
        api_layout.addWidget(api_info)
        
        content_layout.addWidget(api_card)
        
        # Add stretch to push content to the top and let it expand
        content_layout.addStretch(1)
        
        # Add content container to main layout
        main_layout.addWidget(content_container, 1)  # Give it a stretch factor of 1
        
        # Status indicators layout - more compact
        status_layout = QHBoxLayout()
        status_layout.setSpacing(10)
        
        # Folder status indicator
        self.folder_status = QLabel("Pasta: Não selecionada")
        self.folder_status.setStyleSheet("""
            color: #9AA0A6;
            font-size: 11px;
            border: none;
        """)
        status_layout.addWidget(self.folder_status)
        
        # API Key status indicator
        self.api_status = QLabel("API Key: Não configurada")
        self.api_status.setStyleSheet("""
            color: #9AA0A6;
            font-size: 11px;
            border: none;
        """)
        status_layout.addWidget(self.api_status)
        
        status_layout.addStretch()
        
        main_layout.addLayout(status_layout)
        
        # Continue button - wider to match content width
        self.continue_btn = QPushButton("Continuar")
        self.continue_btn.setEnabled(False)
        self.continue_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.continue_btn.setFixedHeight(36)  # Slightly taller for better visibility
        self.continue_btn.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.continue_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                border-radius: 4px;
                font-size: 13px;
                font-weight: bold;
                margin-top: 10px;
                padding: 6px 0;
            }
            QPushButton:hover:enabled {
                background-color: #1E88E5;
            }
            QPushButton:pressed:enabled {
                background-color: #1976D2;
            }
            QPushButton:disabled {
                background-color: #E0E0E0;
                color: #9E9E9E;
            }
        """)
        self.continue_btn.clicked.connect(self.complete_setup)
        
        main_layout.addWidget(self.continue_btn)
        
        # Add stretch to keep everything at the top
        main_layout.addStretch()
        
        self.setLayout(main_layout)
        
        # Connect API key input to validation
        self.api_key_input.textChanged.connect(self.validate_inputs)

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(
            self,
            "Selecionar Pasta de Documentos",
            "",
            QFileDialog.Option.ShowDirsOnly
        )
        if folder:
            self.selected_folder_display.setText(folder)
            self.selected_folder = folder
            self.folder_status.setText(f"Pasta: Selecionada ✓")
            self.folder_status.setStyleSheet("color: #4CAF50; font-size: 13px;")
            self.validate_inputs()

    def validate_inputs(self):
        # Enable continue button only if both folder and API key are provided
        has_folder = hasattr(self, 'selected_folder')
        has_api_key = bool(self.api_key_input.text().strip())
        
        # Update API key status
        if has_api_key:
            self.api_status.setText("API Key: Configurada ✓")
            self.api_status.setStyleSheet("color: #4CAF50; font-size: 13px;")
        else:
            self.api_status.setText("API Key: Não configurada")
            self.api_status.setStyleSheet("color: #9AA0A6; font-size: 13px;")
        
        self.continue_btn.setEnabled(has_folder and has_api_key)

    def complete_setup(self):
        if hasattr(self, 'selected_folder'):
            api_key = self.api_key_input.text().strip()
            # Basic validation for OpenAI API key format (should start with "sk-")
            if not api_key.startswith("sk-"):
                confirm = QMessageBox.question(
                    self,
                    "Formato de Chave Suspeito",
                    "A chave API fornecida não parece estar no formato correto da OpenAI (deve começar com 'sk-').\n\n"
                    "Tem certeza que deseja continuar com esta chave?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                    QMessageBox.StandardButton.No
                )
                
                if confirm == QMessageBox.StandardButton.No:
                    return
            
            # Emit signal with both folder path and API key
            self.setup_complete.emit(self.selected_folder, api_key)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{APP_TITLE} v{APP_VERSION}")
        # Set minimum window size
        self.setMinimumSize(640, 480)  # Reduced from 800x600 to 640x480 (VGA size)
        # Set initial window size to 800x600
        self.resize(800, 600)
        # Center the window on screen
        self.center_on_screen()
        
        # Initialize settings manager
        self.settings_manager = SettingsManager(APP_DATA_DIR)
        self.settings_manager.apply_settings()
        
        # Initialize persistence
        self.persistence = VectorStorePersistence(APP_DATA_DIR)
        
        # Initialize chat history
        self.chat_history = ChatHistory()
        
        # Create central widget and main layout
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)
        
        # Create stacked widget for different views
        self.stacked_widget = QStackedWidget()
        self.main_layout.addWidget(self.stacked_widget)
        
        # Create initial setup view
        self.initial_setup = InitialSetupView()
        self.initial_setup.setup_complete.connect(self.on_setup_complete)
        self.stacked_widget.addWidget(self.initial_setup)
        
        # Create main view
        self.main_view = QWidget()
        self.main_layout_main = QVBoxLayout(self.main_view)
        self.stacked_widget.addWidget(self.main_view)
        
        # Create settings view
        try:
            logger.info("Creating settings view...")
            self.settings_view = SettingsView(self)
            self.stacked_widget.addWidget(self.settings_view)
            logger.info(f"Settings view added to stacked widget, total widgets: {self.stacked_widget.count()}")
        except Exception as e:
            logger.error(f"Error creating settings view: {str(e)}")
            QMessageBox.warning(self, "Aviso", f"Erro ao inicializar configurações: {str(e)}")
        
        # Initialize other components
        self._init_ai_components()
        self._create_ui_components(self.main_layout_main)
        self._create_status_bar()
        self._apply_styles()
        
        # Check if we have existing documents
        has_documents = self._check_for_existing_documents()
        
        if has_documents:
            # If documents exist, skip the folder selection and go straight to main view
            self.stacked_widget.setCurrentWidget(self.main_view)
            self._load_persisted_data()
            # Explicitly ensure buttons are enabled
            QTimer.singleShot(500, self._ensure_buttons_enabled)
            self.status_bar.showMessage("Documentos carregados da base de dados", 3000)
        else:
            # No documents, start with folder selection
            self.stacked_widget.setCurrentWidget(self.initial_setup)
            
        # Start memory usage timer
        self.memory_timer = QTimer()
        self.memory_timer.timeout.connect(self._update_memory_usage)
        self.memory_timer.start(5000)  # Update every 5 seconds

    def _check_for_existing_documents(self):
        """Check if there are already stored documents in the database"""
        try:
            # First try to load vector store from persistence
            success = self.persistence.load_vector_store(self.vector_db)
            
            # Check if we have documents in the vector store
            if success and hasattr(self.vector_db, 'documents') and len(self.vector_db.documents) > 0:
                logger.info(f"Found {len(self.vector_db.documents)} existing documents")
                return True
            
            return False
        except Exception as e:
            logger.error(f"Error checking for existing documents: {str(e)}")
            return False
            
    def on_setup_complete(self, folder_path, api_key):
        """Handle completion of initial setup"""
        try:
            # Store the selected folder path
            self.docs_folder = folder_path
            
            # Save the API key
            if api_key:
                logger.info("Saving API key from initial setup")
                self.settings_manager.set_api_key(api_key)
                # Apply API key to environment immediately
                os.environ["OPENAI_API_KEY"] = api_key
                # Restart AI components with the new API key
                self.restart_ai_components()
            
            # Disable buttons during processing
            self.ask_button.setEnabled(False)
            self.load_docs_button.setEnabled(False)
            
            # Update status
            self.status_bar.showMessage("Iniciando processamento dos documentos...")
            
            # Create progress dialog
            self.progress_dialog = QProgressDialog("Processando documentos...", "Cancelar", 0, 100, self)
            self.progress_dialog.setWindowModality(Qt.WindowModality.WindowModal)
            self.progress_dialog.setWindowTitle("Processando Documentos")
            self.progress_dialog.setMinimumDuration(0)
            self.progress_dialog.setAutoClose(True)
            self.progress_dialog.setAutoReset(True)
            
            # Create and start document loader thread
            self.loader_thread = DocumentLoaderThread(self.vector_db, self.persistence)
            self.loader_thread.progress.connect(self.progress_dialog.setValue)
            self.loader_thread.finished.connect(self._on_documents_loaded)
            self.loader_thread.error.connect(self._on_document_load_error)
            self.loader_thread.file_processed.connect(self._update_document_status)
            self.loader_thread.memory_warning.connect(self._show_memory_warning)
            self.loader_thread.status_update.connect(self.progress_dialog.setLabelText)
            
            # Set the folder path for processing
            self.loader_thread.set_folder_path(self.docs_folder)
            
            # Switch to main view - but keep UI disabled until processing is complete
            self.stacked_widget.setCurrentWidget(self.main_view)
            
            # Start processing
            self.loader_thread.start()
            
            # Connect cancel button
            self.progress_dialog.canceled.connect(self.loader_thread.stop)
            
        except Exception as e:
            logger.error(f"Error in setup completion: {str(e)}")
            QMessageBox.critical(self, "Erro", f"Erro ao iniciar processamento: {str(e)}")
            self.status_bar.showMessage("Erro ao iniciar processamento", 5000)
            self.ask_button.setEnabled(True)
            self.load_docs_button.setEnabled(True)

    def _init_ai_components(self):
        """Initialize AI components"""
        try:
            _ = load_dotenv(find_dotenv())
            
            # Get the API key from settings if we have one
            api_key = self.settings_manager.get_api_key() if hasattr(self, 'settings_manager') else None
            if api_key:
                os.environ["OPENAI_API_KEY"] = api_key
                logger.info("Using API key from settings")
            
            # Initialize the embedder
            self.embedder = OpenAIEmbedder()
            
            # Initialize the vector store
            self.vector_db = SimpleVectorStore(self.embedder)

            self.agent = Agent(
                model=OpenAIChat(id='gpt-4o-mini'),
                knowledge=self.vector_db,
                search_knowledge=True,
                show_tool_calls=False,
                
                markdown=False,
                add_history_to_messages=True,
                instructions='''
                    "Reformulação Interna (não visível ao usuário):
                    Se a pergunta for ampla, ajuste-a para focar em:
                    Palavras-chave (ex: "demonstrações financeiras", "quórum", "deliberação").
                    Datas (ex: "ata de 05/2024").
                    Nomes de sócios/diretores (ex: "voto do Sr. Carlos Martins").
                    Eventos específicos (ex: "aprovação de balanço").
                    Resposta Direta e Objetiva:
                    NUNCA explique seu raciocínio ou métodos de busca.
                    Forneça APENAS:
                    Trechos exatos das atas (se relevantes).
                    Dados como valores, prazos, decisões ou participantes.
                    Referências claras (ex: "página 3, ata de 12/03/2024").
                    Precisão Técnica:
                    Use termos contábeis corretos (ex: "lucro acumulado", "DLPA", "assembléia ordinária").
                    Destaque prazos legais ou obrigações se mencionados na ata.
                    Se a informação não existir:
                    Responda: "Não há registro dessa informação nas atas analisadas."
                    NÃO invente respostas ou use fontes externas.
                    Exemplo de Resposta Ideal:
                    Pergunta: "Quando foi aprovada a distribuição de dividendos em 2023?"
                    Resposta: "15/11/2023, conforme consta na ata de assembleia (página 2). Valor: R$ 200.000,00.""
                '''
            )
            
        except Exception as e:
            logger.error(f"Error initializing components: {str(e)}")
            QMessageBox.critical(self, "Erro de Inicialização", 
                              f"Erro ao inicializar componentes: {str(e)}\n"
                              "Verifique se o arquivo .env está configurado corretamente.")
                              
    def restart_ai_components(self):
        """Restart AI components after settings change"""
        try:
            logger.info("Restarting AI components")
            
            # Apply settings to environment
            if hasattr(self, 'settings_manager'):
                self.settings_manager.apply_settings()
            
            # Validate API key format before continuing
            api_key = self.settings_manager.get_api_key() if hasattr(self, 'settings_manager') else None
            if api_key:
                # Basic validation for OpenAI API key format (should start with "sk-")
                if not api_key.startswith("sk-") or len(api_key) < 20:
                    logger.warning("API key does not appear to be in the correct format")
                    QMessageBox.warning(
                        self,
                        "Chave API Suspeita",
                        "A chave API fornecida não parece estar no formato correto esperado (sk-...).\n\n"
                        "Se você está tendo problemas para fazer perguntas, verifique se a chave está correta.",
                        QMessageBox.StandardButton.Ok
                    )
            
            # Reinitialize the components
            self._init_ai_components()
            
            # Update button states directly
            self._ensure_buttons_enabled()
            
            # Update style to make sure button appearance changes
            if hasattr(self, 'ask_button'):
                self.ask_button.setStyleSheet("")
                QApplication.processEvents()
                # Apply original style from _apply_styles
                self.ask_button.setStyleSheet("""
                    QPushButton {
                        background-color: #5c85d6;
                        color: white;
                        border: none;
                        padding: 10px 18px;
                        border-radius: 6px;
                        font-weight: bold;
                        font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
                    }
                    QPushButton:hover {
                        background-color: #4a6fc3;
                    }
                    QPushButton:pressed {
                        background-color: #3a5fb3;
                    }
                """)
            
            # Show a success message
            self.status_bar.showMessage("Componentes de IA reiniciados com sucesso", 3000)
            
            return True
        except Exception as e:
            logger.error(f"Error restarting AI components: {str(e)}")
            
            # Display user-friendly error
            if "api_key" in str(e).lower() or "authentication" in str(e).lower():
                error_msg = "Erro ao aplicar a chave API. Verifique se a chave é válida e tente novamente."
            else:
                error_msg = f"Erro ao reiniciar componentes: {str(e)}"
                
            QMessageBox.critical(self, "Erro", error_msg)
            return False

    def _create_ui_components(self, layout):
        """Create and arrange UI components"""
        # Define a common font for buttons
        button_font = QFont("Arial", 11)
        
        # History layout - Now first section
        self.history_layout = QVBoxLayout()
        
        # Container for the scrollable history
        history_container = QVBoxLayout()
        
        # Header with controls for history
        history_header = QHBoxLayout()
        
        # History title with modern font
        history_title = QLabel("Histórico de Conversas")
        history_title.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        history_title.setStyleSheet("color: #4a6fc3; margin-bottom: 5px;")
        history_header.addWidget(history_title)
        
        # Add stretch to push buttons to the right
        history_header.addStretch()
        
        # Clear history button with red X emoji
        self.clear_history_button = QPushButton("❌  Limpar Chat")  # Red X emoji
        self.clear_history_button.setToolTip("Limpar Histórico")
        self.clear_history_button.clicked.connect(self.clear_chat_history)
        self.clear_history_button.setFixedSize(160, 32)  # Increased width from 140 to 160
        self.clear_history_button.setFont(QFont("Segoe UI", 11, QFont.Weight.Bold))
        
        self.clear_history_button.setStyleSheet("""
            QPushButton {
                background-color: #ffeded;
                color: #e53e3e;
                border: 1px solid #e53e3e;
                border-radius: 4px;
                padding: 4px 12px;
                margin: 2px;
                text-align: center;
                qproperty-alignment: AlignCenter;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
            }
            QPushButton:hover {
                background-color: #ffe0e0;
                border: 1px solid #c53030;
                color: #c53030;
            }
            QPushButton:pressed {
                background-color: #ffd0d0;
                border: 1px solid #9b2c2c;
                color: #9b2c2c;
            }
        """)
        history_header.addWidget(self.clear_history_button)
        
        # Settings button with icon and modern font
        self.settings_button = QPushButton()
        self.settings_button.setToolTip("Configurações")
        self.settings_button.clicked.connect(self.show_settings)
        # Set fixed size for a smaller button
        self.settings_button.setFixedSize(30, 30)
        # Direct use of icon character in button text
        self.settings_button.setText("⚙")
        # Set smaller font size for better proportions
        self.settings_button.setFont(QFont("Segoe UI", 13))
        # Make icon centered in button
        self.settings_button.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                color: #5c85d6;
                border: none;
                padding: 0;
                margin: 2px;
                text-align: center;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
            }
            QPushButton:hover {
                color: #4a6fc3;
                background-color: rgba(92, 133, 214, 0.1);
            }
            QPushButton:pressed {
                color: #3a5fb3;
            }
        """)
        history_header.addWidget(self.settings_button)
        
        history_container.addLayout(history_header)
        
        # Scroll area for history content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        # Make scroll area take available space and resize with window
        scroll_area.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        
        # Container widget for history content
        self.history_content = QTextEdit()
        self.history_content.setReadOnly(True)
        self.history_content.setAcceptRichText(True)
        
        # Configure history text rendering options
        history_document = self.history_content.document()
        history_document.setDefaultStyleSheet("""
            body {
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
                margin: 0;
                padding: 0;
                font-size: 12px;
                background-color: #f8f9fa;
            }
            .date-header {
                background-color: #eef2f9;
                color: #5c85d6;
                font-weight: bold;
                padding: 10px 14px;
                margin: 20px 0 12px 0;
                border-radius: 6px;
                border-left: 4px solid #5c85d6;
                font-size: 18px;
            }
            .chat-entry {
                margin-bottom: 20px;
                padding: 16px;
                border: 1px solid #e0e4e8;
                border-radius: 8px;
                background-color: white;
            }
            .timestamp {
                color: #939aab;
                font-size: 12pt;
                margin-bottom: 8px;
                text-align: right;
            }
            .question, .answer {
                margin-bottom: 10px;
            }
            .q-label, .a-label {
                font-weight: bold;
                margin-bottom: 4px;
                font-size: 17px;
            }
            .q-label {
                color: #38a169;
            }
            .a-label {
                color: #dd6b20;
            }
            .q-content, .a-content {
                margin-left: 10px;
                line-height: 1.6;
                padding: 6px 0;
                font-size: 16px;
            }
            .q-content {
                background-color: rgba(56, 161, 105, 0.05);
                border-left: 3px solid #38a169;
                padding-left: 12px;
                border-radius: 0 4px 4px 0;
            }
            .a-content {
                background-color: rgba(221, 107, 32, 0.05);
                border-left: 3px solid #dd6b20;
                padding-left: 12px;
                border-radius: 0 4px 4px 0;
            }
        """)
        
        # Set initial empty state
        self.history_content.setHtml("""
        <!DOCTYPE html>
        <html>
        <head></head>
        <body>
            <div style="text-align: center; margin-top: 40px; color: #939aab; font-size: 14px;">
                <div style="margin-bottom: 10px;">✨ Histórico de conversas vazio ✨</div>
                <div>Comece uma nova conversa fazendo uma pergunta!</div>
            </div>
        </body>
        </html>
        """)
        
        # Add to scroll area
        scroll_area.setWidget(self.history_content)
        history_container.addWidget(scroll_area)
        
        self.history_layout.addLayout(history_container)
        
        # Add the history panel to the main layout
        layout.addLayout(self.history_layout)
        
        # Separator line
        separator = QLabel()
        separator.setFixedHeight(1)
        separator.setStyleSheet("background-color: #e0e4e8; margin-top: 5px; margin-bottom: 10px;")
        layout.addWidget(separator)
        
        # Question input with modern font
        question_label = QLabel("Faça sua pergunta sobre as atas:")
        question_label.setFont(QFont("Segoe UI", 14))
        question_label.setStyleSheet("color: #4a6fc3; margin-top: 10px; font-weight: bold;")
        layout.addWidget(question_label)
        
        self.question_input = QTextEdit()
        self.question_input.setPlaceholderText("Carregue documentos para começar...")
        self.question_input.setMaximumHeight(100)
        # Set size policy to make the input area resize properly
        self.question_input.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.question_input.setFont(QFont("Segoe UI", 14))
        self.question_input.setStyleSheet("""
            QTextEdit {
                border: 1px solid #e0e4e8;
                border-radius: 8px;
                padding: 10px;
                background-color: white;
                color: #363a43;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
            }
            QTextEdit:focus {
                border: 1px solid #5c85d6;
                background-color: #fafbfc;
            }
            QTextEdit:disabled {
                background-color: #f0f0f0;
                color: #939aab;
            }
        """)
        # Initially disable the question input until documents are loaded
        self.question_input.setEnabled(False)
        layout.addWidget(self.question_input)
        
        # Buttons with modern font
        button_layout = QHBoxLayout()
        
        self.ask_button = QPushButton("Perguntar")
        self.ask_button.setIcon(QIcon.fromTheme("dialog-question"))
        self.ask_button.clicked.connect(self.process_question)
        self.ask_button.setEnabled(False)  # Disabled until documents are loaded
        self.ask_button.setFont(QFont("Segoe UI", 11))
        self.ask_button.setMinimumHeight(40)
        # Make the button responsive to resizing
        self.ask_button.setSizePolicy(QSizePolicy.Policy.MinimumExpanding, QSizePolicy.Policy.Fixed)
        button_layout.addWidget(self.ask_button)
        
        self.load_docs_button = QPushButton("Carregar Documentos")
        self.load_docs_button.setIcon(QIcon.fromTheme("document-open"))
        self.load_docs_button.clicked.connect(self.load_documents)
        self.load_docs_button.setFont(QFont("Segoe UI", 11))
        self.load_docs_button.setMinimumHeight(40)
        # Make the button responsive to resizing
        self.load_docs_button.setSizePolicy(QSizePolicy.Policy.MinimumExpanding, QSizePolicy.Policy.Fixed)
        button_layout.addWidget(self.load_docs_button)
        
        layout.addLayout(button_layout)
        
        # Hidden response output for storing responses
        # This is needed for the backend but not shown in the UI
        self.response_output = QTextEdit()
        self.response_output.setReadOnly(True)
        self.response_output.setAcceptRichText(True)
        self.response_output.setVisible(False)

    def _create_status_bar(self):
        """Create the status bar"""
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        
        # Create document status label (but not displayed in header)
        self.doc_status_label = QLabel("Sem documentos carregados")
        self.doc_status_label.setStyleSheet("color: #939aab;")
        
        # Add document counter to status bar
        self.doc_count_label = QLabel("Documentos: 0")
        self.status_bar.addPermanentWidget(self.doc_count_label)
        
        # Add chat history counter to status bar
        self.chat_count_label = QLabel("Conversas: 0")
        self.status_bar.addPermanentWidget(self.chat_count_label)
        
        # Add memory usage to status bar
        self.memory_label = QLabel("Memória: 0 MB")
        self.status_bar.addPermanentWidget(self.memory_label)
        
        # Start memory usage timer
        self.memory_timer = QTimer(self)
        self.memory_timer.timeout.connect(self._update_memory_usage)
        self.memory_timer.start(5000)  # Update every 5 seconds
        
    def _update_memory_usage(self):
        """Update memory usage display"""
        try:
            memory_mb = get_memory_usage()
            self.memory_label.setText(f"Memória: {memory_mb:.1f} MB")
            
            # Change color based on memory usage
            if memory_mb > 500:
                self.memory_label.setStyleSheet("color: #e53e3e; font-weight: bold;")  # High memory use (red)
            elif memory_mb > 300:
                self.memory_label.setStyleSheet("color: #dd6b20; font-weight: bold;")  # Medium memory use (orange)
            else:
                self.memory_label.setStyleSheet("color: #515769;")                     # Normal memory use
        except Exception as e:
            logger.error(f"Error updating memory usage: {e}")
            
    def _apply_styles(self):
        """Apply styles to components"""
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f8f9fa;
            }
            QTextEdit {
                border: 1px solid #e0e4e8;
                border-radius: 8px;
                padding: 10px;
                background-color: white;
                color: #363a43;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
            }
            QPushButton {
                background-color: #5c85d6;
                color: white;
                border: none;
                padding: 10px 18px;
                border-radius: 6px;
                font-weight: bold;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
            }
            QPushButton:hover {
                background-color: #4a6fc3;
            }
            QPushButton:pressed {
                background-color: #3a5fb3;
            }
            QPushButton:disabled {
                background-color: #d8dde5;
                color: #939aab;
            }
            QLabel {
                font-weight: bold;
                color: #363a43;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
            }
            QProgressDialog, QMessageBox, QFileDialog {
                background-color: white;
                border-radius: 6px;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
            }
            /* Fix for dialog backgrounds - exclude buttons */
            QDialog {
                background-color: white;
            }
            QDialog QLabel, QDialog QTextEdit, QDialog QLineEdit, QDialog QComboBox, 
            QDialog QSpinBox, QDialog QCheckBox, QDialog QRadioButton {
                background-color: white;
            }
            /* Specific styling for dialog buttons */
            QDialog QPushButton, QMessageBox QPushButton,
            QPushButton[text="OK"], QPushButton[text="Cancelar"], QPushButton[text="Sim"], QPushButton[text="Não"],
            QPushButton[text="Abrir"], QPushButton[text="Salvar"], QPushButton[text="Cancelar"], QPushButton[text="Sim"], 
            QPushButton[text="Não"], QPushButton[text="Abrir"], QPushButton[text="Salvar"] {
                background-color: #5c85d6;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            /* Make sure message box and dialog buttons have distinct colors */
            QMessageBox QPushButton {
                background-color: #5c85d6 !important;
                color: white !important;
            }
            /* Specific styling for Yes/No buttons in confirmation dialogs */
            QMessageBox QPushButton[text="Sim"], QMessageBox QPushButton[text="Não"],
            QMessageBox QPushButton[text="Sim"], QMessageBox QPushButton[text="Não"] {
                min-width: 80px;
                background-color: #5c85d6 !important;
                color: white !important;
            }
            QStatusBar {
                background-color: #e9ecf2;
                color: #515769;
                border-top: 1px solid #d0d4db;
                font-family: 'Segoe UI', 'Roboto', 'Open Sans', sans-serif;
            }
        """)

    def _load_persisted_data(self):
        """Load persisted data from disk"""
        try:
            # Show loading status
            self.status_bar.showMessage("Carregando dados persistidos...")
            QApplication.processEvents()
            
            # Load vector db from persistence
            success = self.persistence.load_vector_store(self.vector_db)
            
            # Load chat history
            chat_history_file = os.path.join(APP_DATA_DIR, "chat_history.json")
            self.chat_history.load_from_disk(chat_history_file)
            
            # Update chat count after loading history
            self._update_chat_count()
            
            # Explicitly update history view to ensure it shows
            self.update_history_view()
            
            # Update UI based on results
            if success:
                # Count unique file sources in metadata instead of total chunks
                sources = set()
                for metadata in self.vector_db.metadatas:
                    if "source" in metadata:
                        sources.add(metadata["source"])
                
                file_count = len(sources)
                chunk_count = len(self.vector_db.documents)
                
                # Document count available in persistence.document_count if loaded
                if file_count > 0:
                    self._update_document_status(f"{file_count} documentos carregados")
                    self.doc_count_label.setText(f"Documentos: {file_count}")
                    
                    # Enable chat functionality
                    self.ask_button.setEnabled(True)
                    self.question_input.setEnabled(True)
                    self.question_input.setPlaceholderText("Digite sua pergunta sobre os documentos aqui...")
                else:
                    self._update_document_status("Nenhum documento encontrado na base de dados")
                    self.ask_button.setEnabled(False)
                    self.question_input.setEnabled(False)
                    
                self.status_bar.showMessage("Dados carregados com sucesso!", 3000)
            else:
                self._update_document_status("Nenhum documento encontrado na base de dados")
                self.status_bar.showMessage("Nenhum dado persistido encontrado", 3000)
                self.ask_button.setEnabled(False)
                self.question_input.setEnabled(False)
                
        except Exception as e:
            logger.error(f"Error loading persisted data: {str(e)}")
            self.status_bar.showMessage("Erro ao carregar dados persistidos", 3000)
            self._update_document_status("Erro ao carregar documentos")
            self.ask_button.setEnabled(False)
            self.question_input.setEnabled(False)

    def _update_document_status(self, text):
        """Update document status display"""
        # Update only the status bar message since the label is not displayed anymore
        self.status_bar.showMessage(text, 3000)

    def process_question(self):
        """Process the user's question"""
        question = self.question_input.toPlainText().strip()
        if not question:
            QMessageBox.warning(self, "Aviso", "Por favor, digite uma pergunta.")
            return
        
        # Temporarily add the question to history with a placeholder response
        timestamp = QDateTime.currentDateTime()
        temp_entry = {
            "timestamp": timestamp,
            "question": question,
            "answer": "⏳ Processando sua pergunta..."
        }
        
        # Add temporary entry to the history
        self.chat_history.history.append(temp_entry)
        
        # Update the history view with the new question immediately
        self.update_history_view()
        
        # Scroll to the bottom of the history view to show the new question
        scrollbar = self.history_content.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
        
        # Disable ask button during processing
        self.ask_button.setEnabled(False)
        self.status_bar.showMessage("Processando pergunta...")
        
        # Clear and set placeholder in the response area
        self.response_output.clear()
        self.response_output.setTextColor(QColor("#939aab"))
        self.response_output.append("⏳ Processando sua pergunta...")
        
        # Create and start worker thread
        self.worker = WorkerThread(self.agent, question)
        self.worker.response_ready.connect(self.handle_response)
        self.worker.error_occurred.connect(self.handle_error)
        self.worker.finished.connect(lambda: self._on_question_finished())
        self.worker.start()
        
        # Clear the question input text
        self.question_input.setPlainText("")

    def _on_question_finished(self):
        """Handle question processing completion"""
        self.ask_button.setEnabled(True)
        self.status_bar.showMessage("Resposta pronta", 3000)

    def handle_response(self, response):
        """Handle the AI response"""
        if not response:
            self.handle_error("Nenhuma resposta recebida do servidor")
            return
            
        try:
            # Log response for debugging
            logger.info(f"Displaying response (length: {len(response)})")
            
            # Find and update the temporary entry with the actual response
            if self.chat_history.history:
                # Get the last entry in history (which should be our temporary entry)
                last_entry = self.chat_history.history[-1]
                if "⏳ Processando sua pergunta..." in last_entry["answer"]:
                    # Replace the placeholder with the actual response
                    self.chat_history.history[-1]["answer"] = response
                else:
                    # If not found with placeholder, use the question from last entry
                    question = last_entry["question"]
                    self.chat_history.add_entry(question, response)
            
            # Update the history view in real-time
            self.update_history_view()
            
            # Clear previous content in response area
            self.response_output.clear()
            
            # Convert the text to HTML with enhanced formatting
            html_content = self._format_response_to_html(response)
            
            # Set the HTML content directly
            self.response_output.setHtml(html_content)
            
            # Move cursor to start
            cursor = self.response_output.textCursor()
            cursor.movePosition(cursor.MoveOperation.Start)
            self.response_output.setTextCursor(cursor)
            
            # Update UI
            QApplication.processEvents()
            
        except Exception as e:
            logger.error(f"Error displaying response: {str(e)}")
            self.handle_error(f"Erro ao exibir resposta: {str(e)}")
            
    def _format_response_to_html(self, text):
        """Convert the response text to rich HTML with enhanced formatting"""
        # First clean any ANSI and formatting artifacts
        clean_text = self._clean_ansi_and_formatting(text.strip())
        
        # Convert markdown to HTML with simpler, more robust formatting
        html = self._simple_text_to_html(clean_text)
        
        # Wrap in a simple HTML document with clean styling
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{
                    font-family: 'Segoe UI', Arial, sans-serif;
                    font-size: 12pt;
                    line-height: 1.5;
                    color: #333;
                    margin: 0;
                    padding: 10px;
                }}
                p {{
                    margin: 0 0 12px 0;
                }}
                h2, h3, h4 {{
                    color: #2a6fc9;
                    margin-top: 16px;
                    margin-bottom: 8px;
                }}
                ul, ol {{
                    margin-bottom: 12px;
                }}
                li {{
                    margin-bottom: 4px;
                }}
                .content {{
                    max-width: 100%;
                }}
            </style>
        </head>
        <body>
            <div class="content">
                {html}
            </div>
        </body>
        </html>
        """
    
    def _simple_text_to_html(self, text):
        """Convert text to HTML with simple, reliable formatting"""
        # Replace special characters
        html = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Split into paragraphs
        paragraphs = html.split('\n\n')
        formatted_paragraphs = []
        
        for paragraph in paragraphs:
            # Skip empty paragraphs
            if not paragraph.strip():
                continue
                
            # Check if this is a bullet list (starts with *, -, or •)
            if paragraph.strip().startswith(('* ', '- ', '• ')):
                # Process as a bullet list
                items = paragraph.split('\n')
                list_html = '<ul>\n'
                for item in items:
                    item = item.strip()
                    if item.startswith(('* ', '- ', '• ')):
                        # Extract the text after the bullet
                        item_text = item[2:].strip()
                        list_html += f'<li>{item_text}</li>\n'
                list_html += '</ul>'
                formatted_paragraphs.append(list_html)
                
            # Check if this is a numbered list (starts with 1., 2., etc.)
            elif paragraph.strip() and paragraph.strip()[0].isdigit() and paragraph.strip().find('. ') > 0:
                # Process as a numbered list
                items = paragraph.split('\n')
                list_html = '<ol>\n'
                for item in items:
                    item = item.strip()
                    if item and item[0].isdigit() and item.find('. ') > 0:
                        # Extract the text after the number
                        item_text = item[item.find('. ')+2:].strip()
                        list_html += f'<li>{item_text}</li>\n'
                list_html += '</ol>'
                formatted_paragraphs.append(list_html)
                
            # Check if this is a heading (starts with # or ##)
            elif paragraph.strip().startswith('#'):
                line = paragraph.strip()
                if line.startswith('# '):
                    heading_text = line[2:]
                    formatted_paragraphs.append(f'<h2>{heading_text}</h2>')
                elif line.startswith('## '):
                    heading_text = line[3:]
                    formatted_paragraphs.append(f'<h3>{heading_text}</h3>')
                elif line.startswith('### '):
                    heading_text = line[4:]
                    formatted_paragraphs.append(f'<h4>{heading_text}</h4>')
                else:
                    # Treat as regular paragraph
                    formatted_paragraphs.append(f'<p>{paragraph.replace("\n", "<br>")}</p>')
            else:
                # Regular paragraph - replace newlines with <br>
                formatted_paragraphs.append(f'<p>{paragraph.replace("\n", "<br>")}</p>')
        
        # Process text formatting within the HTML
        result = '\n'.join(formatted_paragraphs)
        
        # Basic formatting: bold and italic
        # Bold: replace **text** or __text__ with <strong>text</strong>
        result = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', result)
        result = re.sub(r'__(.+?)__', r'<strong>\1</strong>', result)
        
        # Italic: replace *text* or _text_ with <em>text</em>
        result = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', result)
        result = re.sub(r'_([^_]+)_', r'<em>\1</em>', result)
        
        return result
    def _clean_ansi_and_formatting(self, text):
        """Remove ANSI color codes, box drawing characters and other formatting artifacts"""
        import re
        
        if not text:
            return ""
            
        # Remove ANSI escape sequences (color codes)
        ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
        result = ansi_escape.sub('', text)
        
        # Remove box drawing and formatting characters (pipes, etc.)
        result = re.sub(r'[│┌┐└┘┐┘─┬┴┼┤├]', '', result)
        
        # Remove any remaining formatting artifacts
        result = re.sub(r'\[3[0-9]m', '', result)  # Color markers like [34m
        result = re.sub(r'\[0m', '', result)       # Reset markers
        
        # Further cleanup of any artifacts
        return result.strip()

    def handle_error(self, error_message):
        """Handle an error message"""
        logging.error(f"Error occurred: {error_message}")
        
        # Check if this is an OpenAI API key error
        if "401" in error_message and "invalid_api_key" in error_message:
            friendly_message = "Sua chave API OpenAI parece ser inválida. Por favor, acesse as configurações e verifique se a chave foi digitada corretamente."
            self.status_bar.showMessage("Erro: Chave API inválida", 5000)
        else:
            friendly_message = f"Erro: {error_message}"
        
        # Update response area with error
        self.response_output.clear()
        self.response_output.setTextColor(QColor("#e53e3e"))
        self.response_output.append(friendly_message)
        
        # Find and update the temporary entry with the error message
        if self.chat_history.history:
            # Get the last entry in history (which should be our temporary entry)
            last_entry = self.chat_history.history[-1]
            if "⏳ Processando sua pergunta..." in last_entry["answer"]:
                # Replace the placeholder with the error message
                error_html = f'<span style="color: #e53e3e;">{friendly_message}</span>'
                self.chat_history.history[-1]["answer"] = error_html
                
                # Update the history view
                self.update_history_view()
        
        # Reset UI state
        self.ask_button.setEnabled(True)
        self.status_bar.showMessage(friendly_message, 5000)
        
        # If this is an API key error, suggest going to settings
        if "401" in error_message and "invalid_api_key" in error_message:
            QMessageBox.warning(
                self, 
                "Chave API Inválida", 
                "Sua chave API OpenAI parece ser inválida.\n\n"
                "Por favor, acesse as configurações e verifique se a chave foi digitada corretamente.",
                QMessageBox.StandardButton.Ok
            )

    def load_documents(self):
        """Load documents from a selected folder"""
        try:
            # First ask user if they want to reset existing documents
            if hasattr(self, 'vector_db') and self.vector_db and len(self.vector_db.documents) > 0:
                confirm = QMessageBox.question(
                    self,
                    "Resetar Documentos Existentes",
                    "Carregar novos documentos irá substituir todos os documentos atuais. Continuar?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                    QMessageBox.StandardButton.No
                )
                
                if confirm == QMessageBox.StandardButton.No:
                    return
            
            # Let the user select a documents folder
            folder = QFileDialog.getExistingDirectory(
                self,
                "Selecionar Pasta de Documentos",
                "",
                QFileDialog.Option.ShowDirsOnly
            )
            
            if not folder:
                # User canceled the folder selection
                return
                
            # Reset documents before processing the new folder
            logger.info("Resetting existing documents before processing new folder")
            
            # Clear the vector store in memory
            self.vector_db.documents = []
            self.vector_db.vectors = []
            self.vector_db.metadatas = []
            self.vector_db.ids = []
            
            # Delete the vector store files
            if hasattr(self, 'persistence') and self.persistence:
                self.persistence.reset_vector_store()
                logger.info("Vector store files deleted successfully")
            
            # Clear the chat history
            self.chat_history.clear_history()
            
            # Remove history file from disk if it exists
            chat_history_file = os.path.join(APP_DATA_DIR, "chat_history.json")
            try:
                if os.path.exists(chat_history_file):
                    os.remove(chat_history_file)
                    logger.info(f"Removed chat history file: {chat_history_file}")
            except Exception as e:
                logger.error(f"Error removing chat history file: {str(e)}")
            
            # Update the history view
            self.update_history_view()
            
            # Force garbage collection
            gc.collect()
                
            # Store the selected folder path
            self.docs_folder = folder
            
            # Disable buttons during processing
            self.ask_button.setEnabled(False)
            self.load_docs_button.setEnabled(False)
            
            # Update status
            self.status_bar.showMessage("Iniciando processamento dos documentos...")
            
            # Create progress dialog
            self.progress_dialog = QProgressDialog("Processando documentos...", "Cancelar", 0, 100, self)
            self.progress_dialog.setWindowModality(Qt.WindowModality.WindowModal)
            self.progress_dialog.setWindowTitle("Processando Documentos")
            self.progress_dialog.setMinimumDuration(0)
            self.progress_dialog.setAutoClose(True)
            self.progress_dialog.setAutoReset(True)
            
            # Create and start document loader thread
            self.loader_thread = DocumentLoaderThread(self.vector_db, self.persistence)
            self.loader_thread.progress.connect(self.progress_dialog.setValue)
            self.loader_thread.finished.connect(self._on_documents_loaded)
            self.loader_thread.error.connect(self._on_document_load_error)
            self.loader_thread.file_processed.connect(self._update_document_status)
            self.loader_thread.memory_warning.connect(self._show_memory_warning)
            self.loader_thread.status_update.connect(self.progress_dialog.setLabelText)
            
            # Set the folder path for processing
            self.loader_thread.set_folder_path(self.docs_folder)
            
            # Start processing
            self.loader_thread.start()
            
            # Connect cancel button
            self.progress_dialog.canceled.connect(self.loader_thread.stop)
            
        except Exception as e:
            logger.error(f"Error loading documents: {str(e)}")
            QMessageBox.critical(self, "Erro", f"Erro ao carregar documentos: {str(e)}")
            self.status_bar.showMessage("Erro ao carregar documentos", 5000)
            self.ask_button.setEnabled(True)
            self.load_docs_button.setEnabled(True)

    def _show_memory_warning(self, message):
        """Show memory warning message"""
        self.progress_dialog.setLabelText(message)
        self.status_bar.showMessage(message)
        self.memory_label.setText(message)
        self.memory_label.setStyleSheet("color: red; font-weight: bold;")

    def _on_documents_loaded(self):
        """Handle successful document loading"""
        self.progress_dialog.close()
        self.load_docs_button.setEnabled(True)
        
        # Check for API key
        has_api_key = hasattr(self, 'settings_manager') and self.settings_manager.get_api_key() != ""
        
        # Count unique file sources in metadata instead of total chunks
        sources = set()
        for metadata in self.vector_db.metadatas:
            if "source" in metadata:
                sources.add(metadata["source"])
        
        file_count = len(sources)
        chunk_count = len(self.vector_db.documents)
        
        # Update status - show file count, not chunk count
        self._update_document_status(f"{file_count} documentos carregados")
        self.doc_count_label.setText(f"Documentos: {file_count}")
        
        # Show success message
        QMessageBox.information(self, "Sucesso", f"Documentos carregados com sucesso! ({file_count} arquivos, {chunk_count} fragmentos)")
        self.status_bar.showMessage("Documentos carregados com sucesso!", 5000)
        
        # Enable chat functionality only if we have API key
        if has_api_key:
            self.ask_button.setEnabled(True)
            self.question_input.setEnabled(True)
            self.question_input.setPlaceholderText("Digite sua pergunta sobre os documentos aqui...")
        else:
            self.ask_button.setEnabled(False)
            self.question_input.setEnabled(True)
            self.question_input.setPlaceholderText("Configure sua chave API OpenAI nas configurações para fazer perguntas...")
            self.ask_button.setToolTip("Chave API OpenAI não configurada. Acesse as configurações.")
            
            # Show a message about missing API key
            QMessageBox.warning(
                self, 
                "Chave API Necessária", 
                "Seus documentos foram carregados, mas uma chave API OpenAI é necessária para fazer perguntas.\n\n"
                "Por favor, configure sua chave API nas configurações do aplicativo."
            )
        
        # Save the vector store to disk
        self.persistence.save_vector_store(self.vector_db)

    def _on_document_load_error(self, error_message, progress=None):
        """Handle document loading errors"""
        if progress:
            progress.close()
        self.load_docs_button.setEnabled(True)
        self.ask_button.setEnabled(True)
        self._update_document_status("Erro ao carregar documentos")
        QMessageBox.critical(self, "Erro", f"Erro ao carregar documentos: {error_message}")
        self.status_bar.showMessage(f"Erro: {error_message}", 5000)
        
        
    def closeEvent(self, event):
        """Handle the window close event"""
        # Save persistence data
        try:
            self.status_bar.showMessage("Salvando dados...")
            
            # Save vector store
            if hasattr(self, 'vector_db') and self.vector_db and len(self.vector_db.documents) > 0:
                self.persistence.save_vector_store(self.vector_db)
                
            # Save chat history
            if hasattr(self, 'chat_history') and self.chat_history and len(self.chat_history.history) > 0:
                chat_history_file = os.path.join(APP_DATA_DIR, "chat_history.json")
                self.chat_history.save_to_disk(chat_history_file)
                
        except Exception as e:
            logger.error(f"Error saving data on close: {str(e)}")
            QMessageBox.warning(self, "Erro ao salvar", f"Erro ao salvar dados: {str(e)}")
            
        # Accept the close event
        event.accept()
        
    def center_on_screen(self):
        """Center the main window on the screen"""
        # Get the screen geometry
        screen = QApplication.primaryScreen()
        screen_geometry = screen.availableGeometry()
        
        # Get the window size
        window_size = self.frameGeometry()
        
        # Calculate the center point
        center_point = screen_geometry.center()
        
        # Move the window rectangle's center point to the screen's center point
        window_size.moveCenter(center_point)
        
        # Move the window to the top-left point of the centered rectangle
        self.move(window_size.topLeft())

    def update_history_view(self):
        """Update the history view with current chat history"""
        # Save current scroll position
        scrollbar = self.history_content.verticalScrollBar()
        current_position = scrollbar.value()
        
        # Clear current content first
        self.history_content.clear()
        
        # Get formatted history
        formatted_history = self.chat_history.get_formatted_history()
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head></head>
        <body>
            {formatted_history}
        </body>
        </html>
        """
        self.history_content.setHtml(html_content)
        
        # Update the chat count in the status bar
        self._update_chat_count()
        
        # Restore previous scroll position
        scrollbar.setValue(current_position)

    def clear_chat_history(self):
        """Clear the chat history"""
        confirm = QMessageBox.question(
            self,
            "Limpar Histórico",
            "Tem certeza que deseja limpar todo o histórico de conversas?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if confirm == QMessageBox.StandardButton.Yes:
            # Clear the history object
            self.chat_history.clear_history()
            
            # Remove history file from disk if it exists
            chat_history_file = os.path.join(APP_DATA_DIR, "chat_history.json")
            try:
                if os.path.exists(chat_history_file):
                    os.remove(chat_history_file)
                    logger.info(f"Removed chat history file: {chat_history_file}")
            except Exception as e:
                logger.error(f"Error removing chat history file: {str(e)}")
            
            # Set empty state HTML to history content
            self.history_content.clear()
            self.history_content.setHtml("""
            <!DOCTYPE html>
            <html>
            <head></head>
            <body>
                <div style="text-align: center; margin-top: 40px; color: #939aab; font-size: 14px;">
                    <div style="margin-bottom: 10px;">✨ Histórico de conversas vazio ✨</div>
                    <div>Comece uma nova conversa fazendo uma pergunta!</div>
                </div>
            </body>
            </html>
            """)
            
            # Force garbage collection
            gc.collect()
            
            # Update the chat count
            self._update_chat_count()
            
            self.status_bar.showMessage("Histórico limpo com sucesso", 3000)
    
    def _update_chat_count(self):
        """Update the chat count label"""
        count = len(self.chat_history.history)
        self.chat_count_label.setText(f"Conversas: {count}")
        
        # Update the color based on count
        if count > 0:
            self.chat_count_label.setStyleSheet("color: #38a169;")
        else:
            self.chat_count_label.setStyleSheet("color: #515769;")

    def show_settings(self):
        """Show the settings view"""
        try:
            # Log current state
            logger.info(f"Attempting to show settings view, current index: {self.stacked_widget.currentIndex()}")
            logger.info(f"Stacked widget count: {self.stacked_widget.count()}")
            
            # Get widget information for debugging
            for i in range(self.stacked_widget.count()):
                widget = self.stacked_widget.widget(i)
                logger.info(f"Widget at index {i}: {type(widget).__name__}")
            
            # Switch to settings view (should be index 2)
            # If we have 3 widgets, it's at index 2 (0-based indexing)
            if self.stacked_widget.count() >= 3:
                self.stacked_widget.setCurrentIndex(2)
                logger.info(f"Settings view activated, new index: {self.stacked_widget.currentIndex()}")
            else:
                logger.error(f"Settings view not found, widget count: {self.stacked_widget.count()}")
                raise ValueError(f"Settings view not found, widget count: {self.stacked_widget.count()}")
        except Exception as e:
            logger.error(f"Error showing settings: {str(e)}")
            QMessageBox.critical(self, "Erro", f"Não foi possível exibir as configurações: {str(e)}")

    def _ensure_buttons_enabled(self):
        """Force button state update if documents are loaded"""
        has_documents = hasattr(self, 'vector_db') and hasattr(self.vector_db, 'documents') and len(self.vector_db.documents) > 0
        has_api_key = hasattr(self, 'settings_manager') and self.settings_manager.get_api_key() != ""
        
        logger.info(f"Ensuring buttons are enabled: has_documents={has_documents}, has_api_key={has_api_key}")
        
        # Always enable the load_docs_button and settings_button
        self.load_docs_button.setEnabled(True)
        if hasattr(self, 'settings_button'):
            self.settings_button.setEnabled(True)
            
        if has_documents:
            # Only enable asking questions if we have both documents and an API key
            if has_api_key:
                # Update button appearance first
                self.ask_button.setStyleSheet("")
                # Force button to be enabled
                self.ask_button.setEnabled(True)
                self.question_input.setEnabled(True)
                self.question_input.setPlaceholderText("Digite sua pergunta sobre os documentos aqui...")
                self.ask_button.setToolTip("")  # Clear any previous tooltip
                
                # Log that we're enabling the button
                logger.info("Enabling 'Perguntar' button as both documents and API key are available")
                
                # Force GUI update
                QApplication.processEvents()
            else:
                self.ask_button.setEnabled(False)
                self.question_input.setEnabled(True)
                self.question_input.setPlaceholderText("Configure sua chave API OpenAI nas configurações para fazer perguntas...")
                # Give the ask button a different style to indicate it's disabled due to missing API key
                self.ask_button.setToolTip("Chave API OpenAI não configurada. Acesse as configurações.")
                self.ask_button.setStyleSheet("""
                    QPushButton {
                        background-color: #d8dde5;
                        color: #939aab;
                        border: none;
                        padding: 10px 18px;
                        border-radius: 6px;
                        font-weight: bold;
                    }
                """)
        else:
            # No documents loaded
            self.ask_button.setEnabled(False)
            self.question_input.setEnabled(False)
            self.question_input.setPlaceholderText("Carregue documentos para começar...")
            self.ask_button.setToolTip("Carregue documentos para poder fazer perguntas.")

    def reset_documents(self):
        """Reset the document store"""
        confirm = QMessageBox.question(
            self,
            "Resetar Documentos",
            "Tem certeza que deseja resetar todos os documentos carregados?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if confirm == QMessageBox.StandardButton.Yes:
            try:
                # First, delete the vector store files
                if self.persistence.reset_vector_store():
                    logger.info("Vector store files deleted successfully")
                else:
                    logger.warning("Failed to delete some vector store files")
                
                # Clear the vector store in memory
                self.vector_db.documents = []
                self.vector_db.vectors = []
                self.vector_db.metadatas = []
                self.vector_db.ids = []
                
                # Clear the chat history
                self.chat_history.clear_history()
                
                # Remove history file from disk if it exists
                chat_history_file = os.path.join(APP_DATA_DIR, "chat_history.json")
                try:
                    if os.path.exists(chat_history_file):
                        os.remove(chat_history_file)
                        logger.info(f"Removed chat history file: {chat_history_file}")
                except Exception as e:
                    logger.error(f"Error removing chat history file: {str(e)}")
                
                # Set empty state HTML to history content
                self.history_content.clear()
                self.history_content.setHtml("""
                <!DOCTYPE html>
                <html>
                <head></head>
                <body>
                    <div style="text-align: center; margin-top: 40px; color: #939aab; font-size: 14px;">
                        <div style="margin-bottom: 10px;">✨ Histórico de conversas vazio ✨</div>
                        <div>Comece uma nova conversa fazendo uma pergunta!</div>
                    </div>
                </body>
                </html>
                """)
                
                # Update UI elements
                self.ask_button.setEnabled(False)
                self.question_input.setEnabled(False)
                self.question_input.setPlaceholderText("Carregue documentos para começar...")
                self.doc_count_label.setText("Documentos: 0")
                
                # Force garbage collection
                gc.collect()
                
                # Update the chat count
                self._update_chat_count()
                
                # Show success message
                self.status_bar.showMessage("Documentos resetados com sucesso", 3000)
                QMessageBox.information(self, "Sucesso", "Todos os documentos foram resetados com sucesso.")
            
            except Exception as e:
                logger.error(f"Error during document reset: {str(e)}")
                QMessageBox.critical(self, "Erro", f"Erro ao resetar documentos: {str(e)}")
                self.status_bar.showMessage(f"Erro ao resetar documentos: {str(e)}", 5000)

class AgnoCompatibleDocument:
    """Document class compatible with agno framework's expected interface"""
    def __init__(self, content, metadata=None, id=None, score=None):
        self.content = content
        self.metadata = metadata or {}
        self.id = id
        self.score = score
        
    def to_dict(self):
        """Convert to dictionary, which is required by agno"""
        return {
            "content": self.content,
            "metadata": self.metadata,
            "id": self.id,
            "score": self.score
        }
        
    def __repr__(self):
        """String representation of the Document"""
        return f"Document(content={self.content[:50]}..., metadata={self.metadata})"
        
    def __str__(self):
        """String representation for display purposes"""
        source = self.metadata.get('source', 'Desconhecido')
        return f"Documento: {source} (Score: {self.score:.2f if self.score else 'N/A'})"

class ChatHistory:
    """Class to store and manage chat history"""
    def __init__(self, max_entries=50):
        self.history = []
        self.max_entries = max_entries
        
    def add_entry(self, question, answer):
        """Add a question-answer pair to history with timestamp"""
        timestamp = QDateTime.currentDateTime()
        entry = {
            "timestamp": timestamp,
            "question": question,
            "answer": answer
        }
        self.history.append(entry)
        
        # Remove oldest entries if exceeding max_entries
        if len(self.history) > self.max_entries:
            self.history = self.history[-self.max_entries:]
            
    def get_history(self):
        """Return the full history"""
        return self.history
        
    def clear_history(self):
        """Clear all history"""
        # Clear the history list completely
        self.history = []
        
        # Explicitly request garbage collection
        gc.collect()
        
    def get_formatted_history(self):
        """Return history formatted as HTML for display"""
        if not self.history:
            return """
            <div style="text-align: center; margin-top: 40px; color: #939aab; font-size: 14px;">
                <div style="margin-bottom: 10px;">✨ Histórico de conversas vazio ✨</div>
                <div>Comece uma nova conversa fazendo uma pergunta!</div>
            </div>
            """
            
        # Group entries by date
        entries_by_date = {}
        for entry in self.history:
            date_str = entry["timestamp"].toString("yyyy-MM-dd")
            if date_str not in entries_by_date:
                entries_by_date[date_str] = []
            entries_by_date[date_str].append(entry)
        
        html = ""
        # Process entries by date, oldest first
        for date_str in sorted(entries_by_date.keys()):
            # Format the date as a header
            display_date = QDateTime.fromString(date_str, "yyyy-MM-dd").toString("dd 'de' MMMM 'de' yyyy")
            html += f"""
            <div class="date-header">
                {display_date}
            </div>
            """
            
            # Add all conversations for this date, oldest first
            for entry in sorted(entries_by_date[date_str], key=lambda e: e["timestamp"]):
                time_str = entry["timestamp"].toString("hh:mm:ss")
                question = entry["question"].replace('\n', '<br>')
                answer = entry["answer"].replace('\n', '<br>')
                
                html += f"""
                <div class="chat-entry">
                    <div class="timestamp">{time_str}</div>
                    <div class="question">
                        <div class="q-label">Você:</div>
                        <div class="q-content">{question}</div>
                    </div>
                    <div class="answer">
                        <div class="a-label">Resposta:</div>
                        <div class="a-content">{answer}</div>
                    </div>
                </div>
                """
        return html
    
    def save_to_disk(self, file_path):
        """Save chat history to disk"""
        try:
            # Convert QDateTime objects to strings for JSON serialization
            serializable_history = []
            for entry in self.history:
                serializable_entry = {
                    "timestamp": entry["timestamp"].toString("yyyy-MM-ddThh:mm:ss"),
                    "question": entry["question"],
                    "answer": entry["answer"]
                }
                serializable_history.append(serializable_entry)
                
            # Create directory if it doesn't exist
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
                
            # Write to file
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(serializable_history, f, ensure_ascii=False, indent=2)
                
            logger.info(f"Chat history saved to {file_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving chat history: {str(e)}")
            return False
            
    def load_from_disk(self, file_path):
        """Load chat history from disk"""
        try:
            if not os.path.exists(file_path):
                logger.info(f"No chat history file found at {file_path}")
                return False
                
            with open(file_path, 'r', encoding='utf-8') as f:
                serialized_history = json.load(f)
                
            # Convert string timestamps back to QDateTime objects
            self.history = []
            for entry in serialized_history:
                dt = QDateTime.fromString(entry["timestamp"], "yyyy-MM-ddThh:mm:ss")
                self.history.append({
                    "timestamp": dt,
                    "question": entry["question"],
                    "answer": entry["answer"]
                })
                
            logger.info(f"Loaded {len(self.history)} chat history entries from {file_path}")
            return True
        except Exception as e:
            logger.error(f"Error loading chat history: {str(e)}")
            return False

def main():
    """Main application entry point"""
    # Enable high DPI support
    try:
        # Try the newer PyQt6 style first
        QApplication.setAttribute(Qt.ApplicationAttribute.AA_EnableHighDpiScaling, True)
        QApplication.setAttribute(Qt.ApplicationAttribute.AA_UseHighDpiPixmaps, True)
    except AttributeError:
        # Fall back to the older style for PyQt6 6.6.1+
        QApplication.setHighDpiScaleFactorRoundingPolicy(Qt.HighDpiScaleFactorRoundingPolicy.PassThrough)
        # High DPI scaling is enabled by default in newer PyQt versions
    
    # Create application
    app = QApplication(sys.argv)
    app.setApplicationName(APP_TITLE)
    app.setApplicationVersion(APP_VERSION)
    
    # Set application style
    app.setStyle("Fusion")
    
    # Show splash screen
    splash_pix = QPixmap(500, 300)  # Create a 500x300 pixmap for the splash screen
    splash_pix.fill(QColor("#ffffff"))  # Fill with white background

    # Create a painter to draw on the pixmap
    painter = QPainter(splash_pix)
    painter.setRenderHint(QPainter.RenderHint.Antialiasing)
    painter.setRenderHint(QPainter.RenderHint.TextAntialiasing)

    # Create a gradient background
    gradient = QLinearGradient(0, 0, 0, splash_pix.height())
    gradient.setColorAt(0.0, QColor("#4a6fc3"))  # Dark blue at top
    gradient.setColorAt(1.0, QColor("#5c85d6"))  # Lighter blue at bottom
    painter.fillRect(0, 0, splash_pix.width(), splash_pix.height(), gradient)

    # Draw a modern decorative element - a semi-transparent white curve
    path = QPainterPath()
    path.moveTo(0, splash_pix.height() * 0.7)
    path.cubicTo(
        splash_pix.width() * 0.3, splash_pix.height() * 0.9, 
        splash_pix.width() * 0.6, splash_pix.height() * 0.5, 
        splash_pix.width(), splash_pix.height() * 0.8
    )
    path.lineTo(splash_pix.width(), splash_pix.height())
    path.lineTo(0, splash_pix.height())
    path.closeSubpath()
    painter.fillPath(path, QColor(255, 255, 255, 60))  # Semi-transparent white

    # Draw app title
    title_font = QFont("Segoe UI", 24, QFont.Weight.Bold)
    painter.setFont(title_font)
    painter.setPen(QColor("#ffffff"))
    painter.drawText(QRect(0, 70, splash_pix.width(), 50), Qt.AlignmentFlag.AlignCenter, APP_TITLE)

    # Draw app version
    version_font = QFont("Segoe UI", 12)
    painter.setFont(version_font)
    painter.setPen(QColor("#e0e4e8"))
    painter.drawText(QRect(0, 120, splash_pix.width(), 30), Qt.AlignmentFlag.AlignCenter, f"v{APP_VERSION}")

    # Draw loading message
    loading_font = QFont("Segoe UI", 10, QFont.Weight.Normal)
    painter.setFont(loading_font)
    painter.setPen(QColor("#ffffff"))
    painter.drawText(QRect(0, splash_pix.height() - 40, splash_pix.width(), 30), 
                    Qt.AlignmentFlag.AlignCenter, "Inicializando o sistema...")

    # End painting
    painter.end()

    # Create and display the splash screen
    splash = QSplashScreen(splash_pix, Qt.WindowType.WindowStaysOnTopHint)
    splash.setWindowFlags(Qt.WindowType.WindowStaysOnTopHint | Qt.WindowType.FramelessWindowHint)
    splash.show()
    app.processEvents()

    # Add a small delay to show the splash screen longer (500ms)
    start_time = time.time()
    while time.time() - start_time < 0.5:
        app.processEvents()
    
    # Create and show main window
    window = MainWindow()
    window.show()
    
    # Close splash screen
    splash.finish(window)
    
    sys.exit(app.exec())

if __name__ == "__main__":
    main()

